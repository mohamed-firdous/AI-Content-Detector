from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import BinaryIO, List

import fitz  # pymupdf
from docx import Document


class FileProcessingError(Exception):
    """Raised when a file cannot be parsed safely."""


@dataclass
class ParsedDocument:
    filename: str
    raw_text: str
    clean_text: str
    paragraphs: List[str]


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_text(text: str) -> str:
    text = _normalize_whitespace(text)
    text = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)
    return text


def split_into_paragraphs(text: str, min_chars: int = 70) -> List[str]:
    chunks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    merged: List[str] = []
    buffer = ""

    for chunk in chunks:
        if len(chunk) < min_chars:
            buffer = (buffer + " " + chunk).strip()
            continue
        if buffer:
            chunk = f"{buffer} {chunk}".strip()
            buffer = ""
        merged.append(chunk)

    if buffer:
        if merged:
            merged[-1] = f"{merged[-1]} {buffer}".strip()
        else:
            merged.append(buffer)

    return merged


def _read_txt(file_obj: BinaryIO) -> str:
    content = file_obj.read()
    if isinstance(content, bytes):
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise FileProcessingError("Unable to decode TXT file with common encodings.")
    return str(content)


def _read_pdf(file_obj: BinaryIO) -> str:
    try:
        data = file_obj.read()
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:  # pragma: no cover
        raise FileProcessingError(f"Invalid or corrupted PDF: {exc}") from exc

    texts = []
    for page in doc:
        texts.append(page.get_text("text"))
    return "\n".join(texts)


def _read_docx(file_obj: BinaryIO) -> str:
    try:
        data = io.BytesIO(file_obj.read())
        doc = Document(data)
    except Exception as exc:  # pragma: no cover
        raise FileProcessingError(f"Invalid or corrupted DOCX: {exc}") from exc

    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def parse_uploaded_file(filename: str, file_bytes: bytes) -> ParsedDocument:
    ext = filename.lower().split(".")[-1]
    file_obj = io.BytesIO(file_bytes)

    if ext == "txt":
        raw = _read_txt(file_obj)
    elif ext == "pdf":
        raw = _read_pdf(file_obj)
    elif ext == "docx":
        raw = _read_docx(file_obj)
    else:
        raise FileProcessingError("Unsupported file type. Use PDF, DOCX, or TXT.")

    clean = clean_text(raw)
    if not clean:
        raise FileProcessingError("No readable text found in the uploaded file.")

    paragraphs = split_into_paragraphs(clean)
    if not paragraphs:
        raise FileProcessingError("Could not detect meaningful paragraphs in the file.")

    return ParsedDocument(filename=filename, raw_text=raw, clean_text=clean, paragraphs=paragraphs)
